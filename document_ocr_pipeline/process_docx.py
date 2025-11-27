#!/usr/bin/env python3
"""
DOCX 完整处理管道
生成与 PDF 流程一致的输出结构
"""

import sys
import json
import argparse
from pathlib import Path
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import io
from PIL import Image
import subprocess

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from document_ocr_pipeline.extract_document import DocumentExtractor
from document_ocr_pipeline.visualize_extraction import visualize_extraction


def extract_paragraph_content(doc, output_dir, ocr_engine='paddle'):
    """
    提取 Word 文档内容：文本 + 表格 + 图片OCR
    
    Word 文档结构：
    - 段落（Paragraph）
    - 表格（Table）
    - 图片（嵌入在段落或表格中）
    
    策略：
    1. 按顺序遍历文档元素
    2. 提取文本和表格
    3. 提取图片并进行 OCR
    4. 合并高置信度结果
    """
    
    content_data = {
        "paragraphs": [],
        "tables": [],
        "images": [],
        "image_ocr_results": []
    }
    
    # ==================== 阶段 1: 提取文本和表格 ====================
    print(f"\n  📝 阶段1: 提取文档内容（高优先级）...")
    
    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content_data["paragraphs"].append(text)
    
    print(f"    ✓ 段落: {len(content_data['paragraphs'])} 个")
    
    # 提取表格
    for table_idx, table in enumerate(doc.tables, 1):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content_data["tables"].append(table_data)
    
    print(f"    ✓ 表格: {len(content_data['tables'])} 个")
    
    # ==================== 阶段 2: 提取图片并 OCR ====================
    print(f"  🖼️  阶段2: 处理图片内容...")
    
    # 从文档中提取所有图片
    image_count = 0
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                image_count += 1
                image_data = rel.target_part.blob
                
                # 检测图片格式
                img = Image.open(io.BytesIO(image_data))
                img_format = img.format.lower() if img.format else 'png'
                
                image_filename = f"image_{image_count:03d}.{img_format}"
                image_path = output_dir / image_filename
                
                # 保存图片
                with open(image_path, "wb") as f:
                    f.write(image_data)
                
                width, height = img.size
                
                content_data["images"].append({
                    "id": image_count,
                    "path": str(image_path),
                    "format": img_format,
                    "size": [width, height]
                })
                
                print(f"    ✓ 图片 {image_count}: {width}x{height} ({img_format})")
                
                # 对图片运行 OCR
                ocr_json_path = output_dir / f"image_{image_count:03d}_ocr.json"
                try:
                    extractor = DocumentExtractor(ocr_engine=ocr_engine)
                    ocr_result = extractor.extract_from_image(str(image_path))
                    
                    with open(ocr_json_path, 'w', encoding='utf-8') as f:
                        json.dump(ocr_result, f, ensure_ascii=False, indent=2)
                    
                    # 生成可视化
                    vis_path = output_dir / f"image_{image_count:03d}_visualized.png"
                    visualize_extraction(str(image_path), str(ocr_json_path), str(vis_path))
                    
                    content_data["image_ocr_results"].append({
                        "image_id": image_count,
                        "ocr_json": str(ocr_json_path),
                        "visualized": str(vis_path),
                        "text_blocks_count": len(ocr_result.get('text_blocks', []))
                    })
                    
                    print(f"      ✓ OCR: {len(ocr_result.get('text_blocks', []))} 个文本块")
                except Exception as e:
                    print(f"      ✗ OCR失败: {e}")
                    
            except Exception as e:
                print(f"    ✗ 图片 {image_count} 提取失败: {e}")
    
    # ==================== 阶段 3: 合并文本（置信度过滤） ====================
    print(f"  🤖 阶段3: 合并文本内容...")
    
    # 合并段落文本
    direct_text = "\n\n".join(content_data["paragraphs"])
    
    # 合并表格文本
    if content_data["tables"]:
        table_text = "\n\n".join([
            "\n".join([" | ".join(row) for row in table])
            for table in content_data["tables"]
        ])
        direct_text += f"\n\n【表格内容】\n{table_text}"
    
    # 合并高置信度 OCR 文本
    ocr_texts = []
    low_confidence_count = 0
    
    for ocr_result in content_data["image_ocr_results"]:
        ocr_json_path = Path(ocr_result["ocr_json"])
        if ocr_json_path.exists():
            with open(ocr_json_path, 'r', encoding='utf-8') as f:
                ocr_data = json.load(f)
            
            for block in ocr_data.get('text_blocks', []):
                confidence = block.get('confidence', 0.0)
                text = block.get('text', '').strip()
                
                if confidence >= 0.85 and text:
                    ocr_texts.append(text)
                elif text:
                    low_confidence_count += 1
    
    combined_text = direct_text
    if ocr_texts:
        combined_text += f"\n\n【图片文字（高置信度）】\n" + "\n".join(ocr_texts)
    
    if low_confidence_count > 0:
        print(f"    ℹ️  过滤了 {low_confidence_count} 个低置信度文本块")
    
    print(f"    ✓ 合并文本完成: {len(combined_text)} 字符")
    
    # 保存提取结果
    text_json_path = output_dir / "extracted_content.json"
    with open(text_json_path, 'w', encoding='utf-8') as f:
        json.dump({
            "paragraphs": content_data["paragraphs"],
            "tables": content_data["tables"],
            "images": [img["path"] for img in content_data["images"]],
            "combined_text": combined_text
        }, f, ensure_ascii=False, indent=2)
    
    # 生成 VLM 上下文
    vlm_context_path = output_dir / "vlm_context.json"
    vlm_context = {
        "paragraph_count": len(content_data["paragraphs"]),
        "table_count": len(content_data["tables"]),
        "image_count": len(content_data["images"]),
        "ocr_results": content_data["image_ocr_results"]
    }
    
    with open(vlm_context_path, 'w', encoding='utf-8') as f:
        json.dump(vlm_context, f, ensure_ascii=False, indent=2)
    
    # 生成 VLM 提示
    vlm_prompt = f"""# Word 文档内容分析

## 文档结构
- 段落数: {len(content_data['paragraphs'])}
- 表格数: {len(content_data['tables'])}
- 图片数: {len(content_data['images'])}

## 文本内容
{direct_text[:2000]}...

## OCR 提取的图片文字
{chr(10).join(ocr_texts[:10]) if ocr_texts else '无'}

---
请分析文档内容并提取关键信息。
"""
    
    vlm_prompt_path = output_dir / "vlm_prompt.txt"
    with open(vlm_prompt_path, 'w', encoding='utf-8') as f:
        f.write(vlm_prompt)
    
    return {
        "text_combined": combined_text,
        "statistics": {
            "total_paragraphs": len(content_data["paragraphs"]),
            "total_tables": len(content_data["tables"]),
            "total_images": len(content_data["images"]),
            "total_characters": len(combined_text)
        },
        "extracted_content_json": str(text_json_path),
        "vlm_context": str(vlm_context_path),
        "vlm_prompt": str(vlm_prompt_path),
        "images": content_data["images"]
    }


def process_docx(docx_path, output_dir, ocr_engine='paddle'):
    """
    完整处理 DOCX 文件
    生成与 adaptive_ocr_pipeline.py 相同的输出结构
    """
    print(f"🚀 开始处理 DOCX: {docx_path}")
    print(f"📂 输出目录: {output_dir}")
    print(f"🔧 OCR引擎: {ocr_engine}")
    
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_path = Path(docx_path)
    
    # ==================== 步骤 0: 使用 LibreOffice 转换为 PDF 并渲染预览图 ====================
    print(f"\n{'='*70}")
    print(f"📄 步骤 0: 生成页面预览图（LibreOffice 渲染）")
    print(f"{'='*70}")
    
    temp_pdf = output_dir / f"{docx_path.stem}_temp.pdf"
    page_count = 0
    
    try:
        # 调用 LibreOffice 转换 DOCX -> PDF
        print(f"  ⏳ 转换 DOCX 为 PDF...")
        subprocess.run([
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_dir),
            str(docx_path)
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        # LibreOffice 输出的 PDF 文件名与输入文件名相同（仅扩展名不同）
        generated_pdf = output_dir / f"{docx_path.stem}.pdf"
        if generated_pdf.exists() and generated_pdf != temp_pdf:
            generated_pdf.rename(temp_pdf)
        
        print(f"  ✓ PDF 已生成: {temp_pdf.name}")
        
        # 使用 pdfplumber 渲染每一页为图片
        import pdfplumber
        import cv2
        import numpy as np
        
        with pdfplumber.open(temp_pdf) as pdf:
            page_count = len(pdf.pages)
            print(f"  📄 PDF 页数: {page_count}")
            
            for page_num, page in enumerate(pdf.pages, 1):
                # 渲染为高质量图片（300 DPI）
                img = page.to_image(resolution=300)
                img_array = np.array(img.original)
                
                # 保存为 page_XXX_preview.png（与 PDF 流程命名一致）
                preview_path = output_dir / f"page_{page_num:03d}_preview.png"
                cv2.imwrite(str(preview_path), cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR))
                
                height, width = img_array.shape[:2]
                print(f"  ✓ 第 {page_num} 页: {width}x{height}px -> {preview_path.name}")
        
        # 删除临时 PDF 文件
        temp_pdf.unlink()
        print(f"  ✓ 预览图生成完成，临时 PDF 已清理")
        
    except FileNotFoundError:
        print("  ⚠️  警告: 未找到 LibreOffice，跳过预览图生成")
        print("  提示: 安装 LibreOffice 以启用页面预览功能")
        print("  macOS: brew install --cask libreoffice")
    except Exception as e:
        print(f"  ⚠️  预览图生成失败: {e}")
    
    # ==================== 继续原有的内容提取流程 ====================
    doc = Document(str(docx_path))
    
    print(f"\n{'='*70}")
    print(f"📄 处理 Word 文档内容")
    print(f"{'='*70}")
    
    # 提取内容
    extraction_result = extract_paragraph_content(doc, output_dir, ocr_engine)
    
    # ==================== 构建输出结构（模拟 PDF 的 complete_adaptive_ocr.json） ====================
    # Word 文档通常是单页的逻辑结构，但可能有多个物理页面
    # 我们将整个文档作为一个"页面"处理
    
    preview_image = "page_001_preview.png"
    preview_path = output_dir / preview_image
    
    # 如果预览图不存在且有提取的图片，使用第一张图片
    if not preview_path.exists() and extraction_result["images"]:
        first_image_path = Path(extraction_result["images"][0]['path'])
        preview_image = first_image_path.name
    
    # Stage2 OCR 可视化信息
    visualized_image = "page_001_visualized.png"
    visualized_path = output_dir / visualized_image
    
    if preview_path.exists():
        # 直接复制预览图作为可视化结果
        import shutil
        shutil.copy2(preview_path, visualized_path)
    elif extraction_result["images"]:
        # 使用第一张图片的可视化
        first_vis = extraction_result["images"][0]['path'].replace('.', '_visualized.')
        if Path(first_vis).exists():
            import shutil
            shutil.copy2(first_vis, visualized_path)
    
    page_data = {
        "page_number": 1,
        "statistics": extraction_result["statistics"],
        "stage1_global": {
            "image": preview_image,
            "ocr_json": extraction_result["extracted_content_json"],
            "text_source": "direct_extraction"
        },
        "stage2_ocr": {
            "ocr_json": extraction_result["extracted_content_json"],
            "visualized": str(visualized_path) if visualized_path.exists() else ""
        },
        "stage3_vlm": {
            "vlm_prompt": extraction_result["vlm_prompt"],
            "vlm_context": extraction_result["vlm_context"],
            "text_combined": extraction_result["text_combined"]
        }
    }
    
    result = {
        "source_file": str(docx_path),
        "file_type": "docx",
        "total_pages": max(page_count, 1),  # 使用实际 PDF 页数，或至少为 1
        "ocr_engine": ocr_engine,
        "pages": [page_data]
    }
    
    # 保存完整结果
    complete_json = output_dir / "complete_adaptive_ocr.json"
    with open(complete_json, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    
    print(f"\n{'='*70}")
    print(f"✅ 处理完成！")
    print(f"{'='*70}")
    print(f"📊 统计:")
    print(f"  - 物理页数: {page_count}")
    print(f"  - 段落数: {extraction_result['statistics']['total_paragraphs']}")
    print(f"  - 表格数: {extraction_result['statistics']['total_tables']}")
    print(f"  - 图片数: {extraction_result['statistics']['total_images']}")
    print(f"  - 字符数: {extraction_result['statistics']['total_characters']}")
    print(f"  - 输出文件: {complete_json}")
    print(f"  - 输出目录: {output_dir.absolute()}")
    
    return result


def main():
    parser = argparse.ArgumentParser(description='Process DOCX file with OCR and VLM')
    parser.add_argument('docx_file', help='Path to DOCX file')
    parser.add_argument('-o', '--output', help='Output directory', default=None)
    parser.add_argument('--ocr-engine', choices=['easy', 'paddle', 'vision'], 
                       default='paddle', help='OCR engine to use')
    
    args = parser.parse_args()
    
    docx_path = Path(args.docx_file)
    if not docx_path.exists():
        print(f"Error: DOCX file not found: {docx_path}")
        return 1
    
    # 生成输出目录名
    if args.output:
        output_dir = Path(args.output)
    else:
        output_dir = Path(f"{docx_path.stem}_docx_processed")
    
    try:
        process_docx(docx_path, output_dir, args.ocr_engine)
        return 0
    except Exception as e:
        print(f"\n❌ 处理失败: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())




